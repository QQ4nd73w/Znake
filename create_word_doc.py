from docx import Document

doc = Document()

doc.add_heading('3.1 Общая архитектурная схема', level=1)
doc.add_paragraph('Схема: [вставить изображение/ссылку]')
doc.add_paragraph('')
doc.add_paragraph('Краткое пояснение:')
paragraph = doc.add_paragraph()
paragraph.add_run('Проект представляет собой полностью клиентское приложение. ') 
paragraph.add_run('Пользовательский интерфейс реализован в menu.html и index.html, ').bold = False
paragraph.add_run('а вся игровая логика выполняется в script.js. ') 
paragraph.add_run('Данные настроек и рекордов сохраняются в localStorage, поэтому отдельного backend нет. ') 
paragraph.add_run('UI передаёт выбор режима и действия игрока в JavaScript, который управляет игрой и обновляет отображение. ') 
paragraph.add_run('При запуске и завершении сессии script.js записывает и читает данные из localStorage. ')
paragraph.add_run('Внешние API не используются, всё работает в браузере на стороне клиента.')

doc.add_page_break()

doc.add_heading('3.2 Модель данных (ERD)', level=1)
doc.add_paragraph('ERD: [вставить изображение/ссылку]')
doc.add_paragraph('')
doc.add_paragraph('Сущности и назначение:')
doc.add_paragraph('Settings — хранит параметры игры и интерфейса: snakeMode, snakeDifficulty, snakeMuted, snakeSeenMenu, snakeTheme.')
doc.add_paragraph('Records — хранит агрегированные результаты: snakeBest и snakeScores.')
doc.add_paragraph('ScoreEntry — отдельная запись результата с полями score и date.')

doc.add_page_break()

doc.add_heading('3.3 Use Case (Сценарии использования)', level=1)
doc.add_paragraph('Use Case диаграмма: [вставить изображение/ссылку]')
doc.add_paragraph('')
doc.add_paragraph('Ключевые сценарии (happy path):')

table = doc.add_table(rows=1, cols=2)
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Сценарий'
heading_cells[1].text = 'Шаги (кратко)'

rows = [
    ('Выбрать режим и сложность', '1. Открыть menu.html\n2. Выбрать режим и сложность\n3. Нажать «Начать»\n4. Перейти к игровому экрану'),
    ('Начать игру', '1. Инициализировать поле и змейку\n2. Запустить игровой цикл\n3. Отрисовать canvas\n4. Отслеживать ввод игрока'),
    ('Играть и собирать баффы', '1. Управлять змейкой стрелками\n2. Подбирать баффы на поле\n3. Активировать эффекты (x2, магнит, скорость, замедление, bigcoin)\n4. Получать бонусы и менять динамику игры'),
    ('Пауза и продолжение', '1. Нажать пробел\n2. Игра ставится на паузу\n3. Нажать пробел снова\n4. Игра продолжается')
]
for scenario, steps in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = scenario
    row_cells[1].text = steps

doc.save(r'C:\Users\a\Desktop\ПП\Project_Report.docx')
